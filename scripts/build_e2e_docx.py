"""Build E2E_USAGE.docx з вшитими screenshots + clickable TOC.

Usage:
    .venv/bin/python -m scripts.build_e2e_docx

Output:
    docs/E2E_USAGE.docx (~5-10 MB з усіма screenshots)
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

ROOT = Path(__file__).resolve().parent.parent
SHOTS = ROOT / "docs" / "e2e-screenshots"
OUTPUT = ROOT / "docs" / "E2E_USAGE.docx"


# ──────────────────────────────────────────────────────────────────────
# Doc helpers
# ──────────────────────────────────────────────────────────────────────


def add_heading(doc, text: str, level: int) -> None:
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1F, 0x36, 0x49)


def add_para(doc, text: str, *, bold: bool = False, italic: bool = False) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.bold = bold
    run.italic = italic


def add_image(doc, name: str, caption: str) -> None:
    path = SHOTS / name
    if not path.exists():
        add_para(doc, f"[missing screenshot: {name}]", italic=True)
        return
    doc.add_picture(str(path), width=Cm(16))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cr = cap.add_run(f"Figure: {caption}")
    cr.italic = True
    cr.font.size = Pt(9)
    cr.font.color.rgb = RGBColor(0x55, 0x55, 0x55)


def add_code(doc, code: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    for line in code.rstrip("\n").splitlines():
        r = p.add_run(line + "\n")
        r.font.name = "Menlo"
        r.font.size = Pt(9)
        r.font.color.rgb = RGBColor(0x33, 0x33, 0x33)


def add_kv_table(doc, rows: list[tuple[str, str]]) -> None:
    table = doc.add_table(rows=len(rows), cols=2)
    table.style = "Light List Accent 1"
    for i, (k, v) in enumerate(rows):
        c1, c2 = table.rows[i].cells
        c1.text = k
        c2.text = v
    doc.add_paragraph()


def add_clickable_toc(doc) -> None:
    """Insert a Word TOC field. Word renders it as a clickable, hyperlinked
    table-of-contents that updates on `Ctrl+A → F9` or right-click → Update Field.

    Uses standard `TOC \\o "1-3" \\h \\z \\u` field — heading levels 1-3, hyperlinks,
    hide tab leaders, use heading-based outline levels.
    """
    para = doc.add_paragraph()
    run = para.add_run()

    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")

    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = 'TOC \\o "1-3" \\h \\z \\u'

    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")

    placeholder = OxmlElement("w:t")
    placeholder.text = (
        "Right-click and choose 'Update Field' (or press Ctrl+A then F9) "
        "to populate the table of contents."
    )

    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")

    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_sep)
    run._r.append(placeholder)
    run._r.append(fld_end)


# ──────────────────────────────────────────────────────────────────────
# Build
# ──────────────────────────────────────────────────────────────────────


def build() -> Path:
    doc = Document()

    # ── Title page ────────────────────────────────────────────
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = title.add_run("Celmis\nEnd-to-End Usage Guide")
    tr.bold = True
    tr.font.size = Pt(28)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = subtitle.add_run(
        "Multi-repo code Q&A · AI PR review · Cross-repo drift · Repo purge"
    )
    sr.italic = True
    sr.font.size = Pt(13)
    sr.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    mr = meta.add_run("Test run captured 2026-05-25 · Build: Phase 6 + Stage 7/8")
    mr.font.size = Pt(11)
    mr.font.color.rgb = RGBColor(0x77, 0x77, 0x77)

    doc.add_paragraph()
    add_para(doc,
        "This document walks through the full multi-repo Q&A and PR-review "
        "flows end-to-end against three GitLab test repositories "
        "(qa-test-etl, qa-test-chat, qa-test-shared) created specifically for "
        "this scenario. Every screenshot below was captured live during the "
        "run; every metric is a real measurement.")

    doc.add_page_break()

    # ── Clickable TOC ─────────────────────────────────────────
    add_heading(doc, "Table of contents", 1)
    add_para(doc,
        "This table of contents is a live Word field — entries are "
        "clickable hyperlinks once you let Word populate it (right-click → "
        "Update Field, or Ctrl+A then F9).", italic=True)
    add_clickable_toc(doc)
    doc.add_page_break()

    # ── 1. Prerequisites ──────────────────────────────────────
    add_heading(doc, "1. Prerequisites", 1)
    add_para(doc, "Hard requirements on the host machine:")
    for line in [
        "Python 3.13 with a populated .venv (uv sync or pip install -e .[dev])",
        "Node.js 24 + pnpm (Next.js 16 dev server)",
        "Docker (Postgres 17)",
        ".env file at the repo root with POSTGRES_PASSWORD, GEMINI_API_KEY, "
        "QDRANT_URL, QDRANT_API_KEY, NEXTAUTH_SECRET, NEXTAUTH_URL",
    ]:
        doc.add_paragraph(line, style="List Bullet")

    add_para(doc, "Network egress to these hosts is required and enforced "
                  "through src/security/egress.py:")
    for h in ["generativelanguage.googleapis.com (Gemini)",
              "*.cloud.qdrant.io (Qdrant)",
              "gitlab.com, api.github.com, api.bitbucket.org (providers)"]:
        doc.add_paragraph(h, style="List Bullet")

    # ── 2. Starting the stack ─────────────────────────────────
    add_heading(doc, "2. Starting the stack", 1)
    add_code(doc, """\
docker compose up -d postgres
.venv/bin/alembic upgrade head

set -a && source .env && set +a
.venv/bin/python -m uvicorn src.api.main:app --port 8000 --reload   # terminal 1
cd web && pnpm dev                                                  # terminal 2""")

    add_para(doc, "Verify everything is up:")
    add_kv_table(doc, [
        ("Postgres (localhost:5432)", "pg_isready returns 0"),
        ("FastAPI (http://localhost:8000/docs)", "200 (Swagger UI)"),
        ("Next.js (http://localhost:3000)", "307 -> /login when unauthenticated"),
    ])

    # ── 3. Signing in ─────────────────────────────────────────
    add_heading(doc, "3. Signing in", 1)
    add_image(doc, "01-login.png", "NextAuth login page.")
    add_para(doc, "If this is a fresh database, create the first account "
                  "via the API:")
    add_code(doc, """\
curl -s -X POST http://localhost:8000/api/auth/signup \\
  -H "Content-Type: application/json" \\
  -d '{"email":"demo@example.com","password":"DemoPass123!","name":"Demo User"}'""")
    add_image(doc, "02-dashboard.png", "Dashboard immediately after signup.")

    # ── 4. Connecting GitLab ──────────────────────────────────
    add_heading(doc, "4. Connecting GitLab", 1)
    add_para(doc, "Each user stores their own provider tokens, encrypted "
                  "server-side via Fernet (AES-128 + HMAC-SHA256, master key in "
                  "~/code-analysis/secrets/.master.key chmod 600).")
    add_image(doc, "03-connections-empty.png",
              "Connections page before any token is saved.")
    add_para(doc, "Paste a GitLab Personal Access Token (api scope) into "
                  "the GitLab section and click Save & verify.")
    add_image(doc, "04-connections-gitlab-saved.png",
              "GitLab connected; username resolved.")
    add_para(doc, "Implementation detail discovered during this run: GitLab "
                  "/projects?membership=true returns HTTP 500 (~15s) when "
                  "combined with ?order_by=last_activity_at. Fixed by removing "
                  "the offending parameter (src/api/routers/repos.py).",
                  italic=True)

    # ── 5. Adding repositories ────────────────────────────────
    add_heading(doc, "5. Adding repositories", 1)
    add_image(doc, "05-repos-browse-gitlab.png",
              "Browse GitLab repos visible to the connected token.")
    add_para(doc, "Click Add next to each repo. The backend then:")
    for s in [
        "Materialises a repo_slug (e.g. gitlab_acme-qa-test-etl)",
        "Triggers a background clone into ~/code-analysis/repos/{slug}/",
        "Builds the tree-sitter graph at ~/code-analysis/data/{slug}/graph.fdblite",
    ]:
        doc.add_paragraph(s, style="List Bullet")
    add_image(doc, "06-repos-added.png",
              "All three test repos added and indexed.")

    # ── 6. Creating a project ─────────────────────────────────
    add_heading(doc, "6. Creating a project for multi-repo Q&A", 1)
    # 07-projects-list.png and 14-create-project-testgroup.png are not included.
    add_image(doc, "15-project-testgroup-detail.png",
              "Project detail page showing the three sibling repos.")

    # ── 7. Q&A SSE ────────────────────────────────────────────
    add_heading(doc, "7. Asking questions: SSE-streamed Q&A with citations", 1)
    add_para(doc, "Inside the project, hit New chat. The chat UI streams the "
                  "answer in real time over Server-Sent Events.")
    add_para(doc, "For this run we asked a question that requires evidence from "
                  "all three sibling repos:")
    add_para(doc, "Which embedding model do the ETL pipeline and the chat "
                  "service use, where is it configured, and what would break if "
                  "they fall out of sync?", italic=True)
    add_image(doc, "16-chat-multirepo-answer.png",
              "Multi-repo Q&A answer with sources from all three test repos.")
    add_para(doc, "Real numbers from this run:")
    add_kv_table(doc, [
        ("Elapsed", "35.44 s"),
        ("Input tokens", "19 981"),
        ("Output tokens", "2 490"),
        ("Vault hits used", "5 (from all three repos)"),
        ("Source files read", "43"),
    ])
    add_para(doc, "The line numbers cited in the answer are exact — they come "
                  "from the numbered code bundle the retriever sends to the LLM "
                  "(src/qa/multi_repo_retriever.py::_numbered).")

    # ── 8. PR review ──────────────────────────────────────────
    add_heading(doc, "8. Cross-repo PR review: drift detection in action", 1)
    add_para(doc, "The PR-review system is multi-agent (architect, security, "
                  "quality, tests, verifier) and supports GitHub, GitLab, and "
                  "Bitbucket. The differentiator tested here is semantic config "
                  "drift detection — when a PR changes a value other repos in "
                  "the same group depend on but the LLM has no way to 'see' "
                  "the dependency from the diff alone.")
    add_para(doc, "Test scenario", bold=True)
    add_para(doc, "All three repos in qa-test-group share these constants:")
    add_kv_table(doc, [
        ("EMBEDDING_MODEL", "\"gemini-embedding-2\""),
        ("EMBEDDING_DIM", "3072"),
        ("QDRANT_COLLECTION", "\"docs_v1\""),
    ])
    add_para(doc, "We then opened an MR in qa-test-etl that bumps "
                  "EMBEDDING_MODEL to gemini-embedding-3 — a one-line, "
                  "one-file diff.")
    add_code(doc, """\
$ analyzer review "gitlab:acme/qa-test-etl#1" --post

Verdict: CHANGES
  findings:    1 (critical=1)
  agents:      structural, security, tests, quality, architect
  tokens:      3,339/421
  elapsed:     15.4s

Findings:
  Critical src/config.py:29 (architect/arch.cross_repo_drift, conf=1.00)
     Cross-repo drift: Embedding model mismatch with sibling services""")
    add_para(doc, "The drift detector itself took 30 ms to find 9 hardcoded "
                  "locations of 'gemini-embedding-2'. The architect agent "
                  "received that as authoritative input:")
    add_image(doc, "13-gitlab-mr-review.png",
              "GitLab MR with posted critical drift finding + quality warning.")

    # ── 9. Shared mode ────────────────────────────────────────
    add_heading(doc, "9. Why does the demo user see other users' projects?", 1)
    add_para(doc, "Phase 1 of Celmis consciously adopted a shared "
                  "workspace model: all authenticated users see all projects, "
                  "repos, and chats in the workspace. The owner_user_id column "
                  "is stored for audit purposes but is NOT used as a filter in "
                  "read queries:")
    add_code(doc, """\
# src/api/repositories.py:31
async def list_projects(session: AsyncSession) -> list[Project]:
    \"\"\"Усі projects (shared mode — no filter).\"\"\"
    stmt = select(Project).options(selectinload(Project.repos)).order_by(
        Project.updated_at.desc()
    )
    return list((await session.scalars(stmt)).all())""")

    add_para(doc, "Trade-offs", bold=True)
    for s in [
        "Celmis is local-first: it runs on a developer machine or a small "
        "team's VM, not a multi-tenant SaaS. The threat model assumes everyone "
        "with a login is a trusted teammate.",
        "Multi-repo Q&A is much more useful when every team member can ask "
        "about every repo without per-project access management overhead.",
        "Provider tokens (GitLab/GitHub/Bitbucket) are still per-user and "
        "encrypted — only the OWNER of a token can clone with it.",
        "Vault content already passed through a redaction pass "
        "(detect-secrets) when it was indexed.",
    ]:
        doc.add_paragraph(s, style="List Bullet")

    # ── 10. Removing a repository ─────────────────────────────
    add_heading(doc, "10. Removing a repository — the purge flow", 1)
    add_para(doc, "There are two modes for repository removal:")
    add_kv_table(doc, [
        ("Lightweight (default DELETE)",
         "Only this user's auto_review_config row in SQLite. UI hides the "
         "repo for the current user."),
        ("Purge (?purge=true)",
         "Qdrant points, clone, graph, vault notes, Postgres project_repos, "
         "all users' auto_review_config, review_runs, RepoGroup membership."),
    ])

    add_heading(doc, "From the UI", 2)
    add_image(doc, "18-repos-with-purge-tooltip.png",
              "Repositories page with the purge-aware trash icon.")
    add_para(doc, "Clicking the trash icon triggers a two-step confirm:")
    add_code(doc, """\
Step 1:  Remove acme/qa-test-shared?

Step 2:  PURGE COMPLETELY?

         OK = delete clone, graph, vault notes, Qdrant vectors,
              RepoGroup membership, ProjectRepo links, review history.

         Cancel = lightweight remove (only this user's auto-review config).

         Tip: pick "OK" if you want to free disk space and remove the
         repo from Q&A.""")
    add_para(doc, "If you pick OK on step 2, the backend runs the cascade and "
                  "the UI shows a toast like:")
    add_para(doc, "Purged completely — freed 0.0 MB, removed 13 Qdrant points",
             italic=True)
    add_para(doc, "After the purge, the repo disappears from the list:")
    add_image(doc, "19-repos-after-ui-purge.png",
              "Repositories page after UI-triggered purge of qa-test-shared.")

    add_heading(doc, "From the CLI", 2)
    add_code(doc, """\
analyzer repo purge gitlab_acme-qa-test-etl --yes
analyzer repo purge acme-frontend                      # interactive confirm
analyzer repo purge SLUG --skip-qdrant                 # for offline recovery""")

    add_heading(doc, "Real run: purging acme-frontend", 2)
    add_code(doc, """\
$ analyzer repo purge acme-frontend --yes

Purge report:
  Qdrant points removed:   259
  Clone dir removed:       yes (~/code-analysis/repos/acme-frontend)
  Data dir removed:        yes (~/code-analysis/data/acme-frontend)
  Vault dir removed:       yes (~/ObsidianVaults/.../acme-frontend)
  Postgres links removed:  2  (SSE Test Project + Cross-Repo Demo)
  auto_review rows:        1
  review_runs rows:        6
  Disk freed:              39.37 MB
  Elapsed:                 1.73s

Purged 'acme-frontend' completely.""")
    # 17-after-purge removed: same reason.

    add_heading(doc, "Implementation overview", 2)
    add_para(doc, "The cascade lives in src/repos/purge.py (~280 lines). Each "
                  "step is wrapped in its own try/except — a failure in one "
                  "step (e.g. Qdrant unreachable) does NOT block the others. "
                  "The returned PurgeReport has counters per step plus an "
                  "errors list.")
    add_para(doc, "The eight steps (in order):")
    for i, step in enumerate([
        "Count then client.delete Qdrant points with payload.repo == slug",
        "shutil.rmtree the clone directory (with recursive chmod +w pre-pass "
        "to handle git's read-only pack files)",
        "Same for the data directory (graph.fdblite and caches)",
        "Same for the vault directory (PRD/BRD markdown notes)",
        "DELETE FROM project_repos WHERE repo_slug = ? via async SQLAlchemy",
        "DELETE FROM auto_review_config WHERE repo_slug = ? (across all users)",
        "DELETE FROM review_runs WHERE pr_ref LIKE ? "
        "(schema uses pr_ref, not raw slug)",
        "Walk every RepoGroup YAML, remove matching repos via parse_repo_url",
    ], 1):
        doc.add_paragraph(f"{i}. {step}", style="List Number")

    add_para(doc, "API endpoint:")
    add_code(doc, """\
DELETE /api/repos/{slug}              # lightweight; returns 204
DELETE /api/repos/{slug}?purge=true   # cascade; returns PurgeReport JSON""")

    # ── 11. Test repos appendix ───────────────────────────────
    add_heading(doc, "11. Test repos: setup and what they contain", 1)
    add_kv_table(doc, [
        ("qa-test-etl",
         "Python 3.13 ingestion pipeline (30 files) — chunks Markdown, embeds "
         "with gemini-embedding-2, writes to Qdrant docs_v1"),
        ("qa-test-chat",
         "FastAPI 0.135 chat service (69 files) — same embedding model + "
         "same Qdrant collection"),
        ("qa-test-shared",
         "TypeScript 5 shared types (28 files) — EMBEDDING_DIM, "
         "QDRANT_COLLECTION, EmbeddingModel enum"),
    ])
    add_code(doc, """\
# 1. Create empty GitLab projects
for name in qa-test-etl qa-test-chat qa-test-shared; do
  curl -s -X POST -H "PRIVATE-TOKEN: $GITLAB_TOKEN" \\
    -H "Content-Type: application/json" \\
    "https://gitlab.com/api/v4/projects" \\
    -d "{\\"name\\":\\"$name\\",\\"visibility\\":\\"private\\",\\"default_branch\\":\\"main\\"}"
done

# 2. Add them to a group
analyzer group create qa-test-group --desc "Embedding-coherence test group"
for r in qa-test-etl qa-test-chat qa-test-shared; do
  analyzer group add qa-test-group "gitlab:acme/$r"
done

# 3. Index graphs (fast — no LLM calls)
analyzer group index qa-test-group

# 4. Generate vault notes (slow — Gemini-bound, ~18 min for 3 small repos)
analyzer group generate qa-test-group --skip-semgrep""")

    # ── 12. Bugs fixed ────────────────────────────────────────
    add_heading(doc, "12. Bugs discovered and fixed during this E2E", 1)
    bugs = [
        ("analyzer group generate reported 0 modules; all repos failed with "
         "clone_failed pathspec 'dev'",
         "branch=\"dev\" was hardcoded in GroupGenerationOrchestrator "
         "(legacy Bitbucket convention)",
         "Pass branch=None so clone uses the remote's default branch"),
        ("/api/repos/browse/gitlab returned 500 after 15 s",
         "GitLab /projects?membership=true triggers server-side 500 when "
         "combined with order_by=last_activity_at",
         "Drop the order_by param in src/api/routers/repos.py::_browse_gitlab"),
        ("Q&A cited wrong line numbers (off by 5-40)",
         "Code bundle sent to LLM had no line prefixes; LLM guessed positions",
         "Add _numbered() helper that prefixes each line with N; instruct LLM "
         "to use those numbers verbatim"),
        ("Cross-repo drift detector returned 0 matches even when symbol exists",
         "hits.symbols stored as path::symbol; whole-word regex never matched",
         "Normalise path::symbol -> symbol; question idents have priority in "
         "15-target cap"),
        ("analyzer repo purge fails with [Errno 13] Permission denied on src/",
         "Git's working tree sometimes has read-only subdirs (dr-xr-x)",
         "Recursive chmod +w pre-pass before shutil.rmtree "
         "(src/repos/purge.py::_chmod_writable)"),
    ]
    for i, (symptom, cause, fix) in enumerate(bugs, 1):
        add_para(doc, f"Bug {i}: {symptom}", bold=True)
        add_para(doc, f"Cause: {cause}")
        add_para(doc, f"Fix: {fix}")
        doc.add_paragraph()

    # ── 13. Strategic roadmap ─────────────────────────────────
    add_heading(doc, "13. Strategic roadmap", 1)
    add_para(doc, "These are the two largest pieces of new functionality "
                  "identified during the E2E. Both are scoped here for "
                  "execution but not yet implemented.")

    add_heading(doc, "13.1 MCP + skills.md for Claude Code / Forge Code", 2)
    add_para(doc, "Problem.", bold=True)
    add_para(doc, "When a developer asks Claude Code (or any agent harness) "
                  "to add or edit a feature, the agent operates with only "
                  "local file context. For cross-repo monorepo-of-microservices "
                  "setups the agent has no way to know that a change in "
                  "qa-test-etl/src/config.py affects qa-test-chat. Celmis "
                  "already has this index — but it's locked behind a chat UI.")
    add_para(doc, "Solution.", bold=True)
    add_para(doc, "Two layers:")

    add_para(doc, "a) MCP server expansion (src/mcp_server/)", bold=True)
    add_para(doc, "Extend the existing MCP server with cross-repo-aware skills "
                  "the agent calls deterministically during code generation:")
    add_kv_table(doc, [
        ("cross_repo_impact(symbol, repo)",
         "List of {repo, file, line, excerpt} from every sibling repo where "
         "the symbol appears. Use before editing a constant."),
        ("find_similar_features(description)",
         "Top-5 module-level matches across all indexed repos. Prevents "
         "re-implementing existing functionality."),
        ("dependency_graph(file)",
         "Symbol-level callers + callees, plus cross-repo links. Lets the "
         "agent estimate blast radius before refactors."),
        ("proposed_change_review(diff)",
         "Same drift report + architect-style critique that PR review "
         "produces. Agent self-reviews before committing."),
        ("vault_search(question, repos)",
         "The SSE Q&A flow in a single JSON response. Agent uses prose "
         "answers as scratchpad."),
    ])

    add_para(doc, "b) skills.md file at repo root", bold=True)
    add_para(doc, "Claude Code, Forge Code, and similar harnesses look for "
                  "skills.md (or CLAUDE.md / AGENTS.md style files) at the "
                  "repo root. A curated file tells the agent WHEN to reach "
                  "for the MCP tools above. Example:")
    add_code(doc, """\
# Skills available via Celmis MCP

## When editing config files (**/config.py, **/constants.ts, *.env*)
1. Call cross_repo_impact(symbol, current_repo) for each changed constant
2. If sibling repos appear, mention them in the commit message and add a
   "Companion PR needed" checklist to the PR description

## When adding a new feature (new file under src/api/routers/ or src/services/)
1. Call find_similar_features("<one-line description>") first
2. If similar feature exists in another repo, paste its file paths into the
   plan and ask whether to extend the existing implementation instead

## When proposing a refactor (file rename, signature change, class extraction)
1. Call dependency_graph(file) to enumerate callers
2. For each cross-repo caller, generate a separate diff stub
3. Run proposed_change_review(diff) on each before opening PRs""")
    add_para(doc, "Estimated effort: 3-4 days. MCP server extensions are "
                  "mostly thin wrappers around existing retrieval/drift code. "
                  "The skills.md is documentation, ~150 lines.")
    add_para(doc, "Success metric: PR review system should see a measurable "
                  "drop in cross-repo drift findings on PRs authored WITH "
                  "Claude Code — the agent will catch the drift before "
                  "opening the PR.")

    add_heading(doc, "13.2 Product-management feature-proposal system", 2)
    add_para(doc, "Problem.", bold=True)
    add_para(doc, "Today Celmis answers questions reactively. There is no "
                  "proactive surface for 'what should we build next?' — "
                  "questions like 'is this code paying for itself?', 'where "
                  "do users hit friction?', 'what abstraction is begging to "
                  "be extracted?'. Product managers and tech leads currently "
                  "answer these by hand-rolling spreadsheets.")
    add_para(doc, "Solution.", bold=True)
    add_para(doc, "A new src/proposals/ module + UI page that analyses the "
                  "indexed code and surfaces concrete, ranked feature "
                  "proposals.")

    add_para(doc, "Detection inputs (cheap, deterministic):", bold=True)
    add_kv_table(doc, [
        ("Duplicated business logic",
         "Tree-sitter graph + AST similarity → 'Extract validate_email — "
         "found in 5 forms across 3 repos with 90 %+ AST match'"),
        ("N+1 query patterns",
         "Static analysis on SQLAlchemy/Prisma calls → 'Add prefetch in "
         "OrdersController.list — N queries per page load'"),
        ("Missing pagination",
         "Routes that return unbounded lists → 'Add cursor pagination to "
         "/api/chats'"),
        ("A11y gaps",
         "DOM scan of compiled web bundle → '12 <button> without aria-label; "
         "4 forms missing <label>'"),
        ("Slow Q&A queries",
         "Audit log percentiles (elapsed_seconds) → 'Q&A on acme-frontend "
         "averages 31 s — top contributor: 138-file bundle'"),
        ("Cold features",
         "Provider analytics (commits per file × LOC) → 'legacy/captures/ "
         "hasn't been touched in 14 months and has zero callers'"),
        ("Drift-prone configs",
         "Repeat hits in cross-repo drift detector → 'EMBEDDING_MODEL "
         "triggered 3 cross-repo drift findings in 2 weeks'"),
    ])

    add_para(doc, "LLM enrichment.", bold=True)
    add_para(doc, "For each detected pattern, a small Gemini-Flash call "
                  "produces a PRD-style proposal:")
    add_code(doc, """\
title: Extract shared validate_email helper
rationale: |
  Five components (LoginForm.vue, SignupForm.vue, ProfileSettings.vue,
  AdminUserCreate.vue, BulkInviteDialog.vue) implement near-identical
  email validation with subtle differences. A regression in any of them
  is undetectable until production.
ux_impact: |
  Users currently see slightly different error messages on the same
  invalid input across pages. Unifying the helper standardises the
  message and the regex, reducing perceived inconsistency.
scope:
  - acme-frontend/src/composables/validation/useEmail.ts (new)
  - 5 component files (edits)
estimated_effort: 4 hours
risk: low
related_proposals: ["unify-password-strength-checker"]""")

    add_para(doc, "UI surface.", bold=True)
    for s in [
        "A new /proposals page in Next.js with a sortable, filterable table "
        "(severity × category × repos affected)",
        "Lifecycle: proposed -> triaged -> accepted/rejected -> in-progress -> shipped",
        "'Generate PR draft' button that pre-fills a branch with the proposed "
        "diff via the MCP tools from 13.1",
        "Slack/email digest: '5 new proposals this week, 2 high-priority'",
    ]:
        doc.add_paragraph(s, style="List Bullet")

    add_para(doc, "Storage.", bold=True)
    add_kv_table(doc, [
        ("proposals (Postgres)",
         "id, slug, title, rationale, ux_impact, scope_json, "
         "estimated_effort, risk, status, repo_slug, timestamps, "
         "created_by, archived_at"),
        ("proposal_signals",
         "id, proposal_id, signal_type, signal_payload_json, confidence "
         "— backing evidence the user can audit"),
        ("proposal_votes (optional)",
         "proposal_id, user_id, vote (up/down), reason — for small teams "
         "to triage collaboratively"),
    ])
    add_para(doc, "Estimated effort: ~2 weeks for v1 (3 detectors + LLM "
                  "enrichment + simple UI + Postgres schema). ~4 weeks for "
                  "full polish (all 7 detectors + lifecycle workflow + "
                  "Slack integration).")
    add_para(doc, "Success metric: at least one accepted proposal per sprint "
                  "after launch.")

    add_heading(doc, "How the two roadmap items reinforce each other", 2)
    add_para(doc, "The MCP layer (13.1) lets coding agents CONSUME Celmis's "
                  "intelligence at write-time. The proposal system (13.2) lets "
                  "product/eng leads CONSUME it at planning-time. The same "
                  "retrieval and graph infrastructure powers both — proposals "
                  "come with ready-made MCP-compatible diff stubs, agents act "
                  "on them, and the loop closes.")

    # ── 14. Appendix ──────────────────────────────────────────
    add_heading(doc, "14. Appendix: reproducing this document", 1)
    add_para(doc, "Every screenshot in this guide was captured via the Chrome "
                  "DevTools MCP tool against the exact stack described in "
                  "section 2. Expected timing on a 2024 MacBook Pro M3 over "
                  "broadband:")
    add_kv_table(doc, [
        ("Repo creation (3 × GitLab API)", "~2 s"),
        ("group index (3 repos, 127 files)", "35 s"),
        ("group generate (vault, w/ semgrep skipped)", "~18 min (Gemini-bound)"),
        ("Multi-repo Q&A turn (43 files, 20K input tokens)", "35 s"),
        ("Single PR review (1-line diff, 5 agents)", "15 s"),
        ("analyzer repo purge on small repo (~30 files)", "~1 s"),
        ("analyzer repo purge on acme-frontend (~40 MB)", "1.73 s"),
    ])

    # Footer
    foot = doc.add_paragraph()
    foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = foot.add_run(
        "Generated 2026-05-25 — Celmis E2E run on macOS 26 / "
        "Python 3.13 / Node 24"
    )
    fr.italic = True
    fr.font.size = Pt(9)
    fr.font.color.rgb = RGBColor(0x77, 0x77, 0x77)

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    return OUTPUT


if __name__ == "__main__":
    out = build()
    size_kb = out.stat().st_size / 1024
    print(f"OK wrote {out} ({size_kb:.0f} KB)")
